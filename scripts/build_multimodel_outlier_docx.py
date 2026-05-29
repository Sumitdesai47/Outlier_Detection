"""Build layman Word guide for Multimodel Outlier Detection. Run: python scripts/build_multimodel_outlier_docx.py"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Multimodel_Outlier_User_Guide.docx"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading("Multimodel Outlier Detection — simple user guide", level=0)
    doc.add_paragraph(
        "This guide explains how to use the Process Intelligence Dashboard for the "
        "Multimodel Outlier Detection workflow. No technical background is required."
    )

    doc.add_heading("1. Open the portal", level=1)
    for line in (
        "Start the application (your IT team may run: python app.py from the project folder).",
        "Open your web browser and go to: http://127.0.0.1:5001 (port may differ — ask your admin).",
        "In the left sidebar, click Multimodel Outlier Detection.",
    ):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2. Prepare your Excel file", level=1)
    doc.add_paragraph(
        "Your file must be .xlsx (Excel). Use one row per time stamp and one column per measurement (tag)."
    )
    doc.add_paragraph("Required layout (wide format):", style="List Bullet")
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    hdr = ["Timestamp", "Tag_A (example)", "Tag_B (example)", "Tag_C (example)"]
    for i, h in enumerate(hdr):
        table.rows[0].cells[i].text = h
    table.rows[1].cells[0].text = "2026-01-01 00:00:00"
    table.rows[1].cells[1].text = "12.4"
    table.rows[2].cells[0].text = "2026-01-01 01:00:00"
    table.rows[2].cells[1].text = "12.7"
    doc.add_paragraph("Rules:", style="List Bullet")
    for rule in (
        "First column must be named Timestamp (date and time).",
        "All other columns must be numbers only (pressures, temperatures, flows, etc.).",
        "Do not leave empty rows in the middle of the data.",
        "Use the Download sample Excel file button on the upload page to see a correct example.",
    ):
        doc.add_paragraph(rule, style="List Bullet")

    doc.add_heading("3. Upload and configure (step by step)", level=1)
    steps = [
        (
            "Download helpers (optional)",
            "Click Download sample Excel file to see the correct column layout.",
            "Click Download step-by-step guide (this document) to keep instructions offline.",
        ),
        (
            "Upload your file",
            "Click Choose file and select your .xlsx workbook.",
            "Wait until the page says how many tags were loaded — do not click Process until then.",
        ),
        (
            "Tag table — Critical (Crit.)",
            "Check Crit. for tags you want highlighted in summary tables and default settings.",
            "Multimodel training still runs on all numeric tags; Crit. mainly controls which tags you tune and see first in summaries.",
        ),
        (
            "Tag table — Threshold & signals",
            "Threshold: sensitivity (default 3.75 is fine for most users).",
            "Signal checkboxes (S1–S8): leave all on unless an engineer asked you to turn one off.",
        ),
        (
            "Tag table — Plant filter (optional)",
            "Turn on Plant only to drop rows when a tag equals or exceeds a value (e.g. shutdown).",
            "You must enter a number when Plant is on.",
        ),
        (
            "Run analysis",
            "Click Process data. Large files can take several minutes — wait for the results page.",
            "Do not close the browser or click the button twice.",
        ),
    ]
    for i, (title, line1, line2) in enumerate(steps, start=1):
        p = doc.add_paragraph()
        p.add_run(f"Step {i}: ").bold = True
        p.add_run(title).bold = True
        doc.add_paragraph(line1)
        if line2:
            doc.add_paragraph(line2)

    doc.add_heading("4. Read your results", level=1)
    tabs = [
        ("Overview", "High-level counts: tags analyzed, drift events, outliers, health."),
        ("Model Details", "One row per tag: linear/nonlinear path, winning model, features used, flagged points. Click a row for more detail."),
        ("Models & Clusters", "Deep dive for one tag: feature selection stages and all candidate models (CV scores)."),
        ("Tag Analysis", "Search and compare tags by status and anomaly count."),
        ("Graph & Correlation", "Trend chart and related tags for the tag you select."),
        ("Root Cause & Insights", "Plain-language explanation of why points were flagged."),
        ("Event Details", "List of flagged timestamps with explanations."),
        ("Downloads", "Export Excel, CSV, or PDF for reports."),
    ]
    for name, desc in tabs:
        p = doc.add_paragraph()
        p.add_run(name + " — ").bold = True
        p.add_run(desc)

    doc.add_heading("5. Common problems", level=1)
    for item in (
        "“No tags found” — check that row 1 has column names and data starts on row 2.",
        "“Plant filter requires a numeric value” — enter a number or turn Plant off for that tag.",
        "“Select at least one critical tag” — check Crit. on at least one row before Process data.",
        "Page very slow — many tags and long history; try fewer tags or a shorter date range in Excel first.",
        "Server not loading — ask admin to restart the app (scripts/run_flask.ps1).",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("6. What is different in Multimodel mode?", level=1)
    doc.add_paragraph(
        "For each tag, the system builds engineered features, picks linear or nonlinear models using "
        "Pearson and Spearman correlation checks, trains several candidates (e.g. Ridge, ElasticNet, "
        "Gradient Boosting), and uses the best model for S5 peer prediction. You see the winner and "
        "all alternatives on the results page — you do not need to choose the model yourself."
    )

    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
