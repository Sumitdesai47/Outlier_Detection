"""Build Word user guide: Dev (Outlier detection) tab. Run: python scripts/build_dev_outlier_docx.py"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Dev_Outlier_Tab_User_Guide.docx"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    doc.add_heading("Outlier Detection — user guide", level=0)
    doc.add_paragraph(
        "Flask dashboard for time-series outlier workflows. "
        "This document describes the Dev (Outlier detection) tab only."
    )

    doc.add_heading("Run the web app", level=1)
    for line in (
        "Install: pip install -r requirements.txt",
        "Start (from project folder): python app.py",
        "Open: http://127.0.0.1:5001 (override with FLASK_PORT / FLASK_RUN_HOST in .env if needed).",
    ):
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("Dev (Outlier detection) tab", level=1)
    doc.add_paragraph(
        "Open the sidebar link “Dev (Outlier detection)” or use ?tab=part15 on the home page. "
        "Uses the same multi-signal engine as “Multi-signal consensus outlier,” with extra per-tag controls."
    )

    doc.add_heading("Steps (max two lines each)", level=2)
    steps: list[tuple[str, str, str | None]] = [
        (
            "Upload workbook",
            "Choose an .xlsx file with timestamps and numeric tags (wide columns or long format).",
            "The UI loads tag names from the sheet after you select the file.",
        ),
        (
            "Wait for the tag table",
            "After upload, a per-tag grid appears: critical tag, plant filter, threshold, eight signal engines, direction.",
            "Enable “Process data” only after the table is ready and validation is clean.",
        ),
        (
            "Critical tag (Crit.)",
            "Check Crit. for tags you want in the focused results list and for whom threshold / engines / direction apply.",
            "Other tags remain in the dataset but are not controlled by that row’s settings the same way.",
        ),
        (
            "Threshold value",
            "For critical tags, set the reference multiplier vs the preset robust-z scale (default is shown in the table).",
            "Leave the default if you want standard preset behaviour.",
        ),
        (
            "Plant row filter (per tag)",
            "Enable Plant, pick operator and value: rows where (tag operator value) is true are dropped before detection.",
            "OR logic across tags with plant enabled. Dropped rows are not used for limits, training, or plots.",
        ),
        (
            "Signal engines (checkboxes)",
            "Uncheck an engine to skip it for that tag; skipped engines never count as firing.",
            "Defaults follow the multi-signal preset unless you change them.",
        ),
        (
            "Direction",
            "Choose Up, Down, or Both so level-style checks only fire for excursions in that direction for the tag.",
            "Both keeps upward and downward sensitivity.",
        ),
        (
            "Submit",
            "Click Process data; the app runs the pipeline and opens the results page.",
            "Large files can take a while—avoid double-submit.",
        ),
        (
            "Results — summary",
            "The summary lists counts, preset parameters (including S6/S7/S8 when enabled), and applied filters.",
            "Use it to confirm the run matches your form configuration.",
        ),
        (
            "Results — plots and tables",
            "On Dev results, the detail table lists Strong Anomaly rows only; pick a tag for charts.",
            "Anomaly explanation is plain-language plus failed checks; Reason adds metrics and official engine names that fired.",
        ),
        (
            "Download",
            "Use the Excel download on the results page for a full export of the session bundle.",
            "Session data can expire after a long idle period or if the server restarts.",
        ),
    ]

    for i, (title, line1, line2) in enumerate(steps, start=1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(title).bold = True
        doc.add_paragraph(line1)
        if line2:
            doc.add_paragraph(line2)

    doc.add_heading("Optional — without the browser", level=2)
    doc.add_paragraph(
        "CLI: python -m services.dev_outlier_detection_tab <input.xlsx> -o <output.xlsx> "
        "(optional --advanced-json, --critical-tags). Writes Excel sheets without Flask."
    )
    doc.add_paragraph(
        "Streamlit: streamlit run streamlit_app.py (install streamlit). "
        "Separate Dev flow; see services/streamlit_dev_outlier_pipeline.py."
    )

    doc.add_heading(".env configuration", level=2)
    doc.add_paragraph(
        "Set FLASK_SECRET_KEY, FLASK_PORT, and DATABASE_URL if you use database-backed features. "
        "Wrong MySQL credentials do not block upload-only outlier tabs."
    )

    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
