"""Build formatted methodology downloads: Word .docx (python-docx) or Word HTML .doc (stdlib)."""

from __future__ import annotations

import html
import re
from io import BytesIO

_UNDER_EQ = re.compile(r"^=+$")
_UNDER_DASH = re.compile(r"^-{3,}$")
_BULLET = re.compile(r"^-\s+(.+)$")
_NUMBER_DOT = re.compile(r"^(\d+)\.\s+(.+)$")
_TECH_STEP = re.compile(r"^(\d+)\)\s*(.+)$")
_SUB_ITEM = re.compile(r"^([ivxlcdm]+)\)\s*(.+)$", re.I)


def _add_runs_with_markdown(paragraph, text: str) -> None:
    """Support **bold** and `code` spans in a paragraph."""
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def _merge_wrapped_lines(raw_lines: list[str]) -> list[str]:
    """Join soft-wrapped lines (indented continuations that are not new list items)."""
    merged: list[str] = []
    for line in raw_lines:
        if not line.strip():
            merged.append("")
            continue
        indent = len(line) - len(line.lstrip(" "))
        stripped = line.strip()

        is_new_item = (
            bool(_SUB_ITEM.match(stripped))
            or bool(_NUMBER_DOT.match(stripped))
            or bool(_TECH_STEP.match(stripped))
            or stripped.startswith("- ")
        )

        if indent >= 3 and not is_new_item and merged and merged[-1] != "":
            merged[-1] = merged[-1].rstrip() + " " + stripped
        else:
            merged.append(line.rstrip())
    return merged


def _html_inline_markup(text: str) -> str:
    """Escape HTML and render **bold** and `code`."""
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    chunks: list[str] = []
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            chunks.append("<b>" + html.escape(part[2:-2]) + "</b>")
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            chunks.append("<code>" + html.escape(part[1:-1]) + "</code>")
        else:
            chunks.append(html.escape(part))
    return "".join(chunks)


def methodology_text_to_word_html_doc_bytes(plain_body: str) -> bytes:
    """
    Word opens this HTML as a formatted document when the file uses a .doc extension.
    No third-party dependencies (stdlib only).
    """
    raw = plain_body.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    lines = _merge_wrapped_lines(raw)

    parts: list[str] = [
        '<html xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:w="urn:schemas-microsoft-com:office:word">',
        "<head><meta charset=\"utf-8\"><title>Methodology</title>",
        "<style>",
        "body{font-family:Calibri,Arial,sans-serif;font-size:11pt;margin:1in;"
        "line-height:1.35;color:#1e293b;}",
        "h1{font-size:22pt;font-weight:bold;margin:0 0 12pt 0;"
        "padding-bottom:6pt;border-bottom:1px solid #cbd5e1;}",
        "h2{font-size:13pt;font-weight:bold;margin:18pt 0 8pt 0;color:#334155;}",
        "p{margin:6pt 0;}",
        ".li{margin:4pt 0 4pt 18pt;text-indent:-12pt;}",
        ".subli{margin:4pt 0 4pt 36pt;text-indent:-14pt;font-size:10.5pt;}",
        "code{font-family:Consolas,Courier New,monospace;font-size:10pt;}",
        "</style></head><body>",
    ]

    i = 0
    last_blank = True
    while i < len(lines):
        line = lines[i]
        nxt = lines[i + 1] if i + 1 < len(lines) else ""

        if _UNDER_EQ.match(nxt) and line.strip():
            parts.append("<h1>" + html.escape(line.strip()) + "</h1>")
            i += 2
            last_blank = False
            continue

        if _UNDER_DASH.match(nxt) and line.strip():
            parts.append("<h2>" + html.escape(line.strip()) + "</h2>")
            i += 2
            last_blank = False
            continue

        if not line.strip():
            if not last_blank:
                parts.append("<br/>")
            last_blank = True
            i += 1
            continue

        last_blank = False
        stripped = line.strip()

        m_sub = _SUB_ITEM.match(stripped)
        if m_sub:
            inner = f"{m_sub.group(1).lower()}) {m_sub.group(2)}"
            parts.append('<p class="subli">&#8226; ' + _html_inline_markup(inner) + "</p>")
            i += 1
            continue

        m_tech = _TECH_STEP.match(stripped)
        if m_tech:
            inner = f"{m_tech.group(1)}) {m_tech.group(2)}"
            parts.append('<p class="li">' + _html_inline_markup(inner) + "</p>")
            i += 1
            continue

        m_num = _NUMBER_DOT.match(stripped)
        if m_num:
            inner = f"{m_num.group(1)}. {m_num.group(2)}"
            parts.append('<p class="li">' + _html_inline_markup(inner) + "</p>")
            i += 1
            continue

        m_bul = _BULLET.match(line)
        if m_bul:
            parts.append(
                '<p class="li">&#8226; ' + _html_inline_markup(m_bul.group(1)) + "</p>"
            )
            i += 1
            continue

        parts.append("<p>" + _html_inline_markup(stripped) + "</p>")
        i += 1

    parts.append("</body></html>")
    return "".join(parts).encode("utf-8")


def methodology_text_to_docx_bytes(plain_body: str) -> bytes:
    """
    Convert methodology text (sections underlined with --- or ===) to .docx.

    - First line + line of ===  → Title (Heading 0)
    - Section title + line of --- → Heading 1
    - Lines starting with "- " → bullet list
    - Lines like "1. Text" → numbered list
    - Lines like "1) Text" (technical steps) → numbered list
    - Lines like "i) Text" → indented bullet
    """
    from docx import Document
    from docx.shared import Inches, Pt

    raw = plain_body.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    lines = _merge_wrapped_lines(raw)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    i = 0
    last_was_blank = True
    while i < len(lines):
        line = lines[i]
        nxt = lines[i + 1] if i + 1 < len(lines) else ""

        if _UNDER_EQ.match(nxt) and line.strip():
            doc.add_heading(line.strip(), level=0)
            i += 2
            last_was_blank = False
            continue

        if _UNDER_DASH.match(nxt) and line.strip():
            doc.add_heading(line.strip(), level=1)
            i += 2
            last_was_blank = False
            continue

        if not line.strip():
            if not last_was_blank:
                doc.add_paragraph()
            last_was_blank = True
            i += 1
            continue

        last_was_blank = False
        stripped = line.strip()

        m_sub = _SUB_ITEM.match(stripped)
        if m_sub:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.35)
            _add_runs_with_markdown(
                p, f"{m_sub.group(1).lower()}) {m_sub.group(2)}"
            )
            i += 1
            continue

        m_tech = _TECH_STEP.match(stripped)
        if m_tech:
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_markdown(p, f"{m_tech.group(1)}) {m_tech.group(2)}")
            i += 1
            continue

        m_num = _NUMBER_DOT.match(stripped)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_markdown(p, f"{m_num.group(1)}. {m_num.group(2)}")
            i += 1
            continue

        m_bul = _BULLET.match(line)
        if m_bul:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_markdown(p, m_bul.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        _add_runs_with_markdown(p, stripped)
        i += 1

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def build_methodology_download(plain_body: str, base_filename: str) -> tuple[bytes, str, str]:
    """
    Return (payload, filename, mimetype).

    Prefer real .docx when python-docx is available; otherwise Word HTML as .doc
    (still opens in Microsoft Word with headings and lists).
    """
    if not base_filename.lower().endswith(".docx"):
        base = base_filename.rsplit(".", 1)[0] if "." in base_filename else base_filename
    else:
        base = base_filename[:-5]

    try:
        payload = methodology_text_to_docx_bytes(plain_body)
        return (
            payload,
            f"{base}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception:
        pass

    payload = methodology_text_to_word_html_doc_bytes(plain_body)
    return (payload, f"{base}.doc", "application/msword")
